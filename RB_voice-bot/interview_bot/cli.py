# cli.py
import argparse
import os
import PyPDF2
from docx import Document
import pandas as pd
from io import BytesIO
import requests

def parse_arguments():
    parser = argparse.ArgumentParser(description='AI Interview Bot - Enhanced Voice Bot')
    parser.add_argument('--jd', '--job-description', help='Path to job description file (.txt, .pdf, .docx, .csv)')
    parser.add_argument('--resume', '--cv', help='Path to candidate resume file (.txt, .pdf, .docx, .csv)')
    parser.add_argument('--custom-questions', '--questions', help='Path to custom questions file (.txt, .csv) - one question per line')
    parser.add_argument('--voice', choices=['cedar', 'marin', 'alloy', 'echo', 'fable', 'onyx', 'nova', 'shimmer'], help='Voice to use for the interview')
    parser.add_argument('--duration', type=int, default=30, help='Interview duration in minutes (default: 30)')
    return parser.parse_args()



def extract_file_content(file_path):
    """Extract content from various file formats"""
    if not file_path:
        return ""
    
    # Remote URL support
    if isinstance(file_path, str) and (file_path.startswith('http://') or file_path.startswith('https://')):
        try:
            resp = requests.get(file_path, timeout=20)
            resp.raise_for_status()
            content_type = resp.headers.get('Content-Type', '').lower()
            url_ext = os.path.splitext(file_path.split('?')[0])[1].lower()
            data = resp.content
            # Prefer content-type; fallback to URL extension
            if 'application/pdf' in content_type or url_ext == '.pdf':
                reader = PyPDF2.PdfReader(BytesIO(data))
                text = ""
                for page in reader.pages:
                    page_text = page.extract_text() or ""
                    text += page_text + "\n"
                return text
            if 'application/vnd.openxmlformats-officedocument' in content_type or url_ext in ['.docx']:
                doc = Document(BytesIO(data))
                return "\n".join([p.text for p in doc.paragraphs])
            if 'text/plain' in content_type or url_ext in ['.txt', '']:
                try:
                    return resp.text
                except Exception:
                    return data.decode('utf-8', errors='ignore')
            return data.decode('utf-8', errors='ignore')
        except Exception as e:
            raise Exception(f"Error fetching URL {file_path}: {str(e)}")
    
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    
    file_ext = os.path.splitext(file_path)[1].lower()
    
    try:
        if file_ext == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        elif file_ext == '.csv':
            try:
                df = pd.read_csv(file_path)
                return df.to_string()
            except ImportError:
                raise ImportError("pandas required for CSV files. Install: pip install pandas")
        elif file_ext == '.pdf':
            try:
                
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    text = ""
                    for page in reader.pages:
                        text += page.extract_text() + "\n"
                    return text
            except ImportError:
                raise ImportError("PyPDF2 required for PDF files. Install: pip install PyPDF2")
        elif file_ext in ['.docx', '.doc']:
            try:
                
                doc = Document(file_path)
                return "\n".join([paragraph.text for paragraph in doc.paragraphs])
            except ImportError:
                raise ImportError("python-docx required for Word files. Install: pip install python-docx")
        else:
            raise ValueError(f"Unsupported file format: {file_ext}. Supported: .txt, .pdf, .docx, .csv")
    except Exception as e:
        raise Exception(f"Error reading {file_path}: {str(e)}")
